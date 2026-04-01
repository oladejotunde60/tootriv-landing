#!/usr/bin/env python3
"""Generate TooTriv Finance Limited letterhead as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup (A4) ──
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

GREEN_PRIMARY = RGBColor(0x0A, 0x87, 0x54)
GREEN_DARK = RGBColor(0x06, 0x5F, 0x3B)
GOLD = RGBColor(0xF7, 0xB3, 0x2B)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0x9C, 0xA3, 0xAF)
BORDER_COLOR = "0A8754"

# ── Helper functions ──
def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz": 4, "color": "0A8754"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{attrs.get("sz", 4)}" '
            f'w:space="0" w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_run(paragraph, text, size=9, bold=False, color=DARK, font_name='Calibri'):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    return run

def clear_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

# ══════════════════════════════════════════════════════════
# TOP GREEN ACCENT LINE
# ══════════════════════════════════════════════════════════
top_bar = doc.add_paragraph()
top_bar.paragraph_format.space_before = Pt(0)
top_bar.paragraph_format.space_after = Pt(6)
# Use a thin green horizontal rule via bottom border on paragraph
pPr = top_bar._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="18" w:space="1" w:color="{BORDER_COLOR}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# ══════════════════════════════════════════════════════════
# HEADER (2-column table: logo+name | address)
# ══════════════════════════════════════════════════════════
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_table.autofit = True

# Set column widths
for cell in header_table.columns[0].cells:
    cell.width = Cm(9.0)
for cell in header_table.columns[1].cells:
    cell.width = Cm(7.0)

# ── Left cell: Company name ──
left_cell = header_table.cell(0, 0)
clear_cell_borders(left_cell)

p = left_cell.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)

# TooTriv wordmark
run_too = p.add_run("Too")
run_too.font.size = Pt(26)
run_too.font.bold = True
run_too.font.color.rgb = GREEN_PRIMARY
run_too.font.name = 'Calibri'

run_triv = p.add_run("Triv")
run_triv.font.size = Pt(26)
run_triv.font.bold = True
run_triv.font.color.rgb = GOLD
run_triv.font.name = 'Calibri'

# Legal name line
p2 = left_cell.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
run = p2.add_run("FINANCE LIMITED")
run.font.size = Pt(8)
run.font.bold = True
run.font.color.rgb = GRAY
run.font.name = 'Calibri'
run.font.letter_spacing = Pt(2)

# ── Right cell: Address ──
right_cell = header_table.cell(0, 1)
clear_cell_borders(right_cell)

p = right_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(1)

# Registered Office label
add_run(p, "REGISTERED OFFICE", size=7, bold=True, color=GREEN_DARK)
p.add_run("\n")

addr_lines = [
    "2 Ibidapo Shobowale Street",
    "Shofunwa Estate, Ikorodu",
    "Lagos, Nigeria",
]
for line in addr_lines:
    add_run(p, line + "\n", size=8.5, color=GRAY)

p.add_run("\n")
add_run(p, "CONTACT", size=7, bold=True, color=GREEN_DARK)
p.add_run("\n")
add_run(p, "hello@tootriv.com\n", size=8.5, color=GRAY)
add_run(p, "www.tootriv.com", size=8.5, color=GREEN_PRIMARY)

# ── Header separator line ──
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(8)
sep.paragraph_format.space_after = Pt(4)
pPr = sep._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="E6F5EE"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# ══════════════════════════════════════════════════════════
# BODY CONTENT
# ══════════════════════════════════════════════════════════

# Date
date_p = doc.add_paragraph()
date_p.paragraph_format.space_before = Pt(16)
date_p.paragraph_format.space_after = Pt(18)
add_run(date_p, "[Date]", size=10, color=GRAY)

# Recipient
recip = doc.add_paragraph()
recip.paragraph_format.space_before = Pt(0)
recip.paragraph_format.space_after = Pt(18)
recip.paragraph_format.line_spacing = Pt(16)
for line in ["[Recipient Name]", "[Title / Position]", "[Company Name]", "[Address]", "[City, State, Country]"]:
    add_run(recip, line + "\n", size=10, color=DARK)

# Subject
subj = doc.add_paragraph()
subj.paragraph_format.space_before = Pt(0)
subj.paragraph_format.space_after = Pt(16)
add_run(subj, "RE: [SUBJECT LINE]", size=11, bold=True, color=GREEN_DARK)

# Salutation
sal = doc.add_paragraph()
sal.paragraph_format.space_before = Pt(0)
sal.paragraph_format.space_after = Pt(12)
add_run(sal, "Dear [Recipient],", size=10, color=DARK)

# Body paragraphs
body_text = [
    "[Your letter content goes here. This letterhead template is formatted for A4 paper and can be edited directly in Microsoft Word, Google Docs, or LibreOffice.]",
    "[Continue writing your letter content. The formatting, fonts, and colours will be preserved when you edit this document.]",
    "[Add as many paragraphs as needed. The footer will remain at the bottom of the page.]",
]
for text in body_text:
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(10)
    bp.paragraph_format.line_spacing = Pt(16)
    add_run(bp, text, size=10, color=RGBColor(0x37, 0x41, 0x51))

# Closing
closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(24)
closing.paragraph_format.space_after = Pt(4)
add_run(closing, "Yours sincerely,", size=10, color=DARK)

# Signature space
sig = doc.add_paragraph()
sig.paragraph_format.space_before = Pt(0)
sig.paragraph_format.space_after = Pt(0)
add_run(sig, "\n\n", size=10)  # space for signature

# Signature line
sig_line = doc.add_paragraph()
sig_line.paragraph_format.space_before = Pt(0)
sig_line.paragraph_format.space_after = Pt(4)
pPr = sig_line._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="D4D4D4"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)
run = sig_line.add_run("                                        ")  # width placeholder
run.font.size = Pt(10)

# Name and title
name_p = doc.add_paragraph()
name_p.paragraph_format.space_before = Pt(4)
name_p.paragraph_format.space_after = Pt(2)
add_run(name_p, "Olusegun Tunde", size=11, bold=True, color=GREEN_DARK)

title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(0)
add_run(title_p, "Founder & CEO, TooTriv Finance Limited", size=9, color=GRAY)

# ══════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════

# Add spacing before footer
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(40)
spacer.paragraph_format.space_after = Pt(4)

# Footer separator
footer_sep = doc.add_paragraph()
footer_sep.paragraph_format.space_before = Pt(0)
footer_sep.paragraph_format.space_after = Pt(6)
pPr = footer_sep._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="E6F5EE"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# Footer table (3 columns)
footer_table = doc.add_table(rows=1, cols=3)
footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for cell in footer_table.columns[0].cells:
    cell.width = Cm(5.0)
for cell in footer_table.columns[1].cells:
    cell.width = Cm(6.0)
for cell in footer_table.columns[2].cells:
    cell.width = Cm(5.0)

# Left: copyright
fl = footer_table.cell(0, 0)
clear_cell_borders(fl)
p = fl.paragraphs[0]
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
add_run(p, "© 2026 TooTriv Finance Limited\n", size=7, color=LIGHT_GRAY)
add_run(p, "RC: [Registration Number]", size=7, color=LIGHT_GRAY)

# Center: tagline
fc = footer_table.cell(0, 1)
clear_cell_borders(fc)
p = fc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('"Instant, affordable salary advances\nfor every Nigerian worker"')
run.font.size = Pt(7)
run.font.italic = True
run.font.color.rgb = GREEN_PRIMARY
run.font.name = 'Calibri'

# Right: website
fr = footer_table.cell(0, 2)
clear_cell_borders(fr)
p = fr.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
add_run(p, "www.tootriv.com\n", size=7, bold=True, color=GREEN_PRIMARY)
add_run(p, "hello@tootriv.com", size=7, color=LIGHT_GRAY)

# Bottom green accent
bottom_bar = doc.add_paragraph()
bottom_bar.paragraph_format.space_before = Pt(6)
bottom_bar.paragraph_format.space_after = Pt(0)
pPr = bottom_bar._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{BORDER_COLOR}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# ── Save ──
output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, "TooTriv_Letterhead.docx")
doc.save(output_path)
print(f"✅ Letterhead saved to: {output_path}")
